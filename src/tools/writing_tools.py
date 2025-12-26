"""Writing tools for book generation and style analysis"""

import logging
import os
import json
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime

logger = logging.getLogger('writing_tools')

from ..tools.book_retrieval_tools import get_book_metadata, get_full_book_text, search_by_author, search_by_genre
from ..tools.vector_tools import search_vectors
from ..tools.embedding_tools import create_embeddings

logger = logging.getLogger('writing_tools')


def analyze_author_style(book_ids: List[str], author: Optional[str] = None) -> Dict[str, Any]:
    """
    Extract writing style from author's books
    
    Args:
        book_ids: List of book IDs to analyze
        author: Optional author name for additional search
        
    Returns:
        Dictionary with style characteristics
    """
    style_data = {
        'author': author,
        'books_analyzed': len(book_ids),
        'average_sentence_length': 0,
        'vocabulary_complexity': 0,
        'dialogue_ratio': 0,
        'description_ratio': 0,
        'common_phrases': [],
        'writing_patterns': []
    }
    
    all_text = ""
    
    for book_id in book_ids:
        book_text = get_full_book_text(book_id)
        if book_text:
            all_text += book_text + "\n\n"
    
    if not all_text:
        return style_data
    
    # Analyze text
    sentences = all_text.split('.')
    words = all_text.split()
    
    # Average sentence length
    if sentences:
        style_data['average_sentence_length'] = sum(len(s.split()) for s in sentences) / len(sentences)
    
    # Vocabulary complexity (unique words / total words)
    unique_words = len(set(word.lower() for word in words))
    style_data['vocabulary_complexity'] = unique_words / len(words) if words else 0
    
    # Dialogue ratio (lines with quotes)
    dialogue_lines = sum(1 for line in all_text.split('\n') if '"' in line or "'" in line)
    total_lines = len(all_text.split('\n'))
    style_data['dialogue_ratio'] = dialogue_lines / total_lines if total_lines > 0 else 0
    
    # Description ratio (approximate - lines with descriptive words)
    descriptive_words = ['was', 'were', 'seemed', 'appeared', 'looked', 'felt', 'saw', 'heard']
    description_lines = sum(1 for line in all_text.split('\n') 
                           if any(word in line.lower() for word in descriptive_words))
    style_data['description_ratio'] = description_lines / total_lines if total_lines > 0 else 0
    
    return style_data


def analyze_genre_style(genre: str, top_k: int = 10) -> Dict[str, Any]:
    """
    Extract genre-specific writing patterns
    
    Args:
        genre: Genre name
        top_k: Number of books to analyze
        
    Returns:
        Dictionary with genre style characteristics
    """
    # Search for books in genre
    genre_books = search_by_genre(genre, top_k)
    
    book_ids = [book.get('metadata', {}).get('book_id') for book in genre_books 
                if book.get('metadata', {}).get('book_id')]
    
    # Analyze style from genre books
    style_data = analyze_author_style(book_ids)
    style_data['genre'] = genre
    style_data['books_analyzed'] = len(book_ids)
    
    return style_data


def extract_story_structure(book_id: str) -> Dict[str, Any]:
    """
    Extract plot structure from a book
    
    Args:
        book_id: Book ID
        
    Returns:
        Dictionary with story structure information
    """
    book_text = get_full_book_text(book_id)
    metadata = get_book_metadata(book_id)
    
    if not book_text:
        return {}
    
    structure = {
        'book_id': book_id,
        'title': metadata.get('title', ''),
        'chapters': metadata.get('chapters', []),
        'estimated_acts': [],
        'key_events': []
    }
    
    # Simple structure detection
    chapters = metadata.get('chapters', [])
    total_chapters = len(chapters)
    
    if total_chapters > 0:
        # Estimate 3-act structure
        act1_end = total_chapters // 3
        act2_end = (total_chapters * 2) // 3
        
        structure['estimated_acts'] = [
            {'act': 1, 'chapters': list(range(1, act1_end + 1))},
            {'act': 2, 'chapters': list(range(act1_end + 1, act2_end + 1))},
            {'act': 3, 'chapters': list(range(act2_end + 1, total_chapters + 1))}
        ]
    
    return structure


def generate_character_sheet(character_data: Dict[str, Any]) -> str:
    """
    Create character development sheet
    
    Args:
        character_data: Character information dictionary
        
    Returns:
        Markdown formatted character sheet
    """
    sheet = f"# Character: {character_data.get('name', 'Unknown')}\n\n"
    
    if character_data.get('description'):
        sheet += f"## Description\n\n{character_data.get('description')}\n\n"
    
    if character_data.get('traits'):
        sheet += "## Traits\n\n"
        for trait in character_data.get('traits', []):
            sheet += f"- {trait}\n"
        sheet += "\n"
    
    if character_data.get('background'):
        sheet += f"## Background\n\n{character_data.get('background')}\n\n"
    
    if character_data.get('relationships'):
        sheet += "## Relationships\n\n"
        for rel in character_data.get('relationships', []):
            sheet += f"- {rel}\n"
        sheet += "\n"
    
    return sheet


def track_continuity(project_id: str, new_content: str, consistency_db: Dict[str, Any]) -> Dict[str, Any]:
    """
    Check consistency with existing content
    
    Args:
        project_id: Project ID
        new_content: New content to check
        consistency_db: Consistency database
        
    Returns:
        Dictionary with consistency check results
    """
    issues = []
    warnings = []
    
    project_data = consistency_db.get(project_id, {})
    characters = project_data.get('characters', {})
    locations = project_data.get('locations', {})
    plot_points = project_data.get('plot_points', [])
    
    # Check character consistency (simple name matching)
    content_lower = new_content.lower()
    for char_name, char_data in characters.items():
        if char_name.lower() in content_lower:
            # Character mentioned - check if details match
            if char_data.get('description'):
                # Future: Add sophisticated character detail checking
                # For now, just track that character was mentioned
                pass
    
    # Check location consistency
    for loc_name, loc_data in locations.items():
        if loc_name.lower() in content_lower:
            # Location mentioned - future: verify location details match
            pass
    
    return {
        'project_id': project_id,
        'issues': issues,
        'warnings': warnings,
        'consistent': len(issues) == 0
    }


def refine_style(text: str, target_style: Dict[str, Any], model_config: Optional[Dict[str, Any]] = None) -> str:
    """
    Adjust writing style to match target using LLM
    
    Args:
        text: Text to refine
        target_style: Target style characteristics
        model_config: Optional model configuration (uses default if not provided)
        
    Returns:
        Refined text matching target style
    """
    try:
        from ..tools.llm_tools import call_llm
        from ..tools.model_manager import get_model_manager
        
        # Get model for refinement
        if not model_config:
            model_manager = get_model_manager()
            model_config = model_manager.get_model_for_task("refinement")
        
        if not model_config:
            logger.warning("No refinement model configured, returning original text")
            return text
        
        # Build style description from target_style
        style_desc = []
        if target_style.get('average_sentence_length'):
            style_desc.append(f"Average sentence length: {target_style['average_sentence_length']:.1f} words")
        if target_style.get('vocabulary_complexity'):
            style_desc.append(f"Vocabulary complexity: {target_style['vocabulary_complexity']:.2f}")
        if target_style.get('dialogue_ratio'):
            style_desc.append(f"Dialogue ratio: {target_style['dialogue_ratio']:.2f}")
        if target_style.get('description_ratio'):
            style_desc.append(f"Description ratio: {target_style['description_ratio']:.2f}")
        
        style_guide = "\n".join(style_desc) if style_desc else "Match the general writing style"
        
        prompt = f"""Refine the following text to match this writing style:

Style characteristics:
{style_guide}

Original text:
{text[:2000]}

Provide the refined version that matches the target style:"""
        
        model_name = model_config.get('model', 'mistral:7b')
        provider = model_config.get('provider', 'ollama')
        
        refined_text = call_llm(
            prompt=prompt,
            model=model_name,
            provider=provider,
            max_tokens=len(text.split()) * 2
        )
        
        return refined_text if refined_text else text
        
    except Exception as e:
        logger.error(f"Error refining style: {e}", exc_info=True)
        return text  # Return original on error


def export_to_markdown(content: str, metadata: Dict[str, Any], output_path: str) -> bool:
    """
    Export content to Markdown file
    
    Args:
        content: Book content
        metadata: Book metadata
        output_path: Output file path
        
    Returns:
        True if successful
    """
    try:
        path_obj = Path(output_path)
        path_obj.parent.mkdir(parents=True, exist_ok=True)
        
        with open(path_obj, 'w', encoding='utf-8') as f:
            # Write frontmatter
            f.write("---\n")
            for key, value in metadata.items():
                f.write(f"{key}: {value}\n")
            f.write("---\n\n")
            
            # Write content
            f.write(content)
        
        return True
    except Exception as e:
        logger.error(f"Error exporting to Markdown: {e}", exc_info=True)
        return False


def export_to_docx(content: str, metadata: Dict[str, Any], output_path: str) -> bool:
    """
    Export content to Word document
    
    Args:
        content: Book content
        metadata: Book metadata
        output_path: Output file path
        
    Returns:
        True if successful
    """
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(metadata.get('title', 'Untitled'), 0)
        
        # Add metadata
        if metadata.get('author'):
            doc.add_paragraph(f"Author: {metadata.get('author')}")
        if metadata.get('series'):
            doc.add_paragraph(f"Series: {metadata.get('series')}")
        
        doc.add_paragraph()  # Blank line
        
        # Add content
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(output_path)
        return True
    except Exception as e:
        logger.error(f"Error exporting to DOCX: {e}", exc_info=True)
        return False


def export_to_epub(content: str, metadata: Dict[str, Any], output_path: str) -> bool:
    """
    Export content to EPUB file
    
    Args:
        content: Book content
        metadata: Book metadata
        output_path: Output file path
        
    Returns:
        True if successful
    """
    try:
        import ebooklib
        from ebooklib import epub
        
        book = epub.EpubBook()
        
        # Set metadata
        book.set_identifier(metadata.get('book_id', str(datetime.now().timestamp())))
        book.set_title(metadata.get('title', 'Untitled'))
        book.set_language('en')
        
        if metadata.get('author'):
            book.add_author(metadata.get('author'))
        
        # Create chapter
        chapter = epub.EpubHtml(
            title=metadata.get('title', 'Chapter 1'),
            file_name='chap_01.xhtml',
            lang='en'
        )
        
        # Convert content to HTML
        content_html = content.replace('\n\n', '</p><p>')
        content_html = f'<p>{content_html}</p>'
        chapter.content = content_html.encode('utf-8')
        
        book.add_item(chapter)
        
        # Add to spine
        book.spine = [chapter]
        
        # Add default NCX and Nav file
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Write EPUB
        epub.write_epub(output_path, book)
        return True
    except Exception as e:
        logger.error(f"Error exporting to EPUB: {e}", exc_info=True)
        return False

